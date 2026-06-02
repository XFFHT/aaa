from __future__ import annotations

import argparse
import base64
import json
import mimetypes
import os
import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any
from uuid import uuid4
from xml.sax.saxutils import escape

import fitz
import requests
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from pptx import Presentation
from pptx.dml.color import RGBColor
from pptx.enum.text import PP_ALIGN
from pptx.util import Inches, Pt as PptPt


ROOT = Path(__file__).resolve().parent
DEFAULT_WORK_DIR = ROOT / "work"
DEFAULT_OUTPUT_DIR = ROOT / "outputs"
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}
TEXT_EXTS = {".txt", ".md"}


def load_env(path: Path = ROOT / ".env") -> None:
    if not path.exists():
        return
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        value = value.strip().strip('"').strip("'")
        os.environ.setdefault(key.strip(), value)


load_env()


def env_path(name: str, default: Path) -> Path:
    return Path(os.getenv(name, str(default))).expanduser()


def now_id() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid4().hex[:8]


def mkdir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def compress_image(src: Path, dst: Path, max_side: int = 1800) -> Path:
    with Image.open(src) as image:
        image = image.convert("RGB")
        image.thumbnail((max_side, max_side))
        image.save(dst, "JPEG", quality=82, optimize=True)
    return dst


def prepare_source(source_path: Path, job_dir: Path, max_pages: int) -> dict[str, Any]:
    ext = source_path.suffix.lower()
    images_dir = mkdir(job_dir / "pages")
    images: list[Path] = []
    text = ""

    if ext == ".pdf":
        with fitz.open(str(source_path)) as pdf:
            extracted: list[str] = []
            page_count = min(len(pdf), max_pages)
            for index in range(page_count):
                page = pdf[index]
                page_text = page.get_text("text").strip()
                if page_text:
                    extracted.append(f"[第 {index + 1} 页文字]\n{page_text}")
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                png_path = images_dir / f"page_{index + 1:03d}.png"
                jpg_path = images_dir / f"page_{index + 1:03d}.jpg"
                pix.save(str(png_path))
                compress_image(png_path, jpg_path)
                png_path.unlink(missing_ok=True)
                images.append(jpg_path)
            text = "\n\n".join(extracted)
        return {"kind": "pdf", "images": images, "text": text}

    if ext in IMAGE_EXTS:
        jpg_path = images_dir / "page_001.jpg"
        compress_image(source_path, jpg_path)
        images.append(jpg_path)
        return {"kind": "image", "images": images, "text": ""}

    if ext == ".docx":
        return {"kind": "docx", "images": [], "text": read_docx_text(source_path)}

    if ext in TEXT_EXTS:
        return {"kind": "text", "images": [], "text": source_path.read_text(encoding="utf-8", errors="ignore")}

    raise ValueError(f"暂不支持这个文件类型: {source_path.suffix}")


def image_data_url(path: Path) -> str:
    mime = mimetypes.guess_type(path.name)[0] or "image/jpeg"
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def llm_prompt(title: str, subject: str, grade: str, source_kind: str) -> str:
    return f"""
你是一个给老师减负的试卷加工助手。请把老师上传的旧试卷、扫描件或图片加工成可交付教学资产包。

目标：
1. 只提取印刷体试题内容，忽略学生手写答案、批改痕迹、涂鸦、圈画、红笔批注、页眉页脚广告。
2. 自动形成“去字迹清版”的干净试卷版本。
3. 识别每道题的题型、知识点、答案、解析。
4. 判断难度与适配度，难度用 1-5 分，适配度用 1-5 分。
5. 为每道题生成 1-2 道同知识点、同难度或略变式题。
6. 输出严格 JSON，不要 Markdown，不要解释文字。

用户给出的基本信息：
- 标题：{title or "未提供"}
- 学科：{subject or "未提供"}
- 年级：{grade or "未提供"}
- 来源类型：{source_kind}

JSON 结构必须是：
{{
  "title": "试卷标题",
  "subject": "学科",
  "grade": "年级",
  "source_quality": "来源质量简评",
  "processing_notes": ["处理说明"],
  "questions": [
    {{
      "number": "题号",
      "type": "题型",
      "stem": "原题题干，已经去除手写痕迹",
      "options": ["A. ...", "B. ..."],
      "clean_version": "清版后的完整题目文本",
      "answer": "答案",
      "explanation": "解析",
      "knowledge_points": ["知识点1", "知识点2"],
      "difficulty": {{"level": 3, "label": "中等", "reason": "原因"}},
      "suitability": {{"grade": "适合年级", "stage": "适合阶段", "score": 4, "reason": "原因"}},
      "variants": [
        {{
          "stem": "变式题题干",
          "options": ["A. ..."],
          "answer": "答案",
          "explanation": "解析",
          "knowledge_points": ["知识点"],
          "difficulty_level": 3
        }}
      ]
    }}
  ]
}}

如果有题目看不清，请仍尽力复原，并在 processing_notes 或对应题目 explanation 中标记“不确定”。
""".strip()


def post_chat_completion(payload: dict[str, Any]) -> dict[str, Any]:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    base_url = os.getenv("OPENAI_BASE_URL", "").strip().rstrip("/")
    if not api_key or not base_url:
        raise RuntimeError("OPENAI_API_KEY 或 OPENAI_BASE_URL 没有配置")

    url = f"{base_url}/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    session = requests.Session()
    session.trust_env = os.getenv("LLM_TRUST_ENV_PROXY", "0").strip().lower() in {"1", "true", "yes", "on"}
    response = session.post(url, headers=headers, json=payload, timeout=900)

    if response.status_code >= 400 and "response_format" in payload:
        fallback = dict(payload)
        fallback.pop("response_format", None)
        response = session.post(url, headers=headers, json=fallback, timeout=900)

    try:
        data = response.json()
    except Exception:
        data = {"text": response.text[:3000]}
    if response.status_code >= 400:
        raise RuntimeError(f"LLM HTTP {response.status_code}: {data}")
    return data


def extract_json(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start >= 0 and end > start:
            return json.loads(cleaned[start : end + 1])
        raise


def analyze_with_llm(source: dict[str, Any], title: str, subject: str, grade: str) -> dict[str, Any]:
    model = os.getenv("OPENAI_MODEL", "gpt-5.5").strip()
    content: list[dict[str, Any]] = [
        {
            "type": "text",
            "text": llm_prompt(title=title, subject=subject, grade=grade, source_kind=source["kind"]),
        }
    ]
    content.append(
        {
            "type": "text",
            "text": (
                "Additional required workflow: select the high-value questions from the old paper and use them "
                "to compose a new paper. For each question add fields: selected_for_new_paper (boolean), "
                "selection_reason, teaching_focus, and lesson_steps. The variants field should contain similar "
                "practice questions for the teaching PPT. Prefer quality over quantity; if the source has few "
                "questions, select all clear and teachable questions."
            ),
        }
    )
    if source.get("text"):
        content.append({"type": "text", "text": "以下是可提取文字，图片优先，文字作为辅助：\n" + source["text"][:60000]})
    for image in source.get("images") or []:
        content.append({"type": "image_url", "image_url": {"url": image_data_url(Path(image))}})

    payload: dict[str, Any] = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是严谨的中文教育资源结构化处理助手，只输出合法 JSON。"},
            {"role": "user", "content": content},
        ],
        "temperature": 0.15,
        "response_format": {"type": "json_object"},
    }
    data = post_chat_completion(payload)
    message = data.get("choices", [{}])[0].get("message", {})
    text = message.get("content", "")
    if not text:
        raise RuntimeError(f"LLM 返回为空: {data}")
    return extract_json(text)


def as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, str):
        return [value] if value.strip() else []
    return [value]


def clean_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    text = str(value).strip()
    return text or default


def normalize_package(raw: dict[str, Any], title: str, subject: str, grade: str) -> dict[str, Any]:
    questions: list[dict[str, Any]] = []
    for index, item in enumerate(as_list(raw.get("questions")), start=1):
        if not isinstance(item, dict):
            item = {"stem": str(item)}
        difficulty = item.get("difficulty") if isinstance(item.get("difficulty"), dict) else {}
        suitability = item.get("suitability") if isinstance(item.get("suitability"), dict) else {}
        options = [clean_text(x) for x in as_list(item.get("options")) if clean_text(x)]
        variants = []
        for variant in as_list(item.get("variants")):
            if not isinstance(variant, dict):
                variant = {"stem": str(variant)}
            variants.append(
                {
                    "stem": clean_text(variant.get("stem")),
                    "options": [clean_text(x) for x in as_list(variant.get("options")) if clean_text(x)],
                    "answer": clean_text(variant.get("answer")),
                    "explanation": clean_text(variant.get("explanation")),
                    "knowledge_points": [clean_text(x) for x in as_list(variant.get("knowledge_points")) if clean_text(x)],
                    "difficulty_level": int(variant.get("difficulty_level") or difficulty.get("level") or 3),
                }
            )
        questions.append(
            {
                "number": clean_text(item.get("number"), str(index)),
                "type": clean_text(item.get("type"), "未识别"),
                "stem": clean_text(item.get("stem") or item.get("clean_version")),
                "options": options,
                "clean_version": clean_text(item.get("clean_version") or item.get("stem")),
                "answer": clean_text(item.get("answer")),
                "explanation": clean_text(item.get("explanation")),
                "knowledge_points": [clean_text(x) for x in as_list(item.get("knowledge_points")) if clean_text(x)],
                "difficulty": {
                    "level": int(difficulty.get("level") or 3),
                    "label": clean_text(difficulty.get("label"), "中等"),
                    "reason": clean_text(difficulty.get("reason")),
                },
                "suitability": {
                    "grade": clean_text(suitability.get("grade"), grade or "未标注"),
                    "stage": clean_text(suitability.get("stage")),
                    "score": int(suitability.get("score") or 3),
                    "reason": clean_text(suitability.get("reason")),
                },
                "variants": variants,
            }
        )

    return {
        "title": clean_text(raw.get("title"), title or "旧试卷教学资产包"),
        "subject": clean_text(raw.get("subject"), subject or "未标注"),
        "grade": clean_text(raw.get("grade"), grade or "未标注"),
        "source_quality": clean_text(raw.get("source_quality"), "已完成自动识别"),
        "processing_notes": [clean_text(x) for x in as_list(raw.get("processing_notes")) if clean_text(x)],
        "questions": questions,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def setup_docx_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_run_font(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_docx_paragraph(document: Document, text: str = "", style: str | None = None):
    para = document.add_paragraph(text, style=style)
    set_run_font(para)
    return para


def build_docx(package: dict[str, Any], path: Path) -> None:
    doc = Document()
    setup_docx_font(doc)

    doc.add_heading(package["title"], level=0)
    add_docx_paragraph(doc, f"学科：{package['subject']}    年级：{package['grade']}    生成时间：{package['generated_at']}")
    add_docx_paragraph(doc, f"来源质量：{package['source_quality']}")
    if package["processing_notes"]:
        add_docx_paragraph(doc, "处理说明：" + "；".join(package["processing_notes"]))

    doc.add_heading("一、清版试卷", level=1)
    for q in package["questions"]:
        add_docx_paragraph(doc, f"{q['number']}. {q['clean_version'] or q['stem']}")
        for option in q["options"]:
            add_docx_paragraph(doc, option)

    doc.add_page_break()
    doc.add_heading("二、答案与解析", level=1)
    for q in package["questions"]:
        add_docx_paragraph(doc, f"{q['number']}. 答案：{q['answer'] or '未识别'}")
        if q["explanation"]:
            add_docx_paragraph(doc, f"解析：{q['explanation']}")

    doc.add_heading("三、题型、知识点、难度与适配度", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["题号", "题型", "知识点", "难度", "适配度", "说明"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for q in package["questions"]:
        cells = table.add_row().cells
        cells[0].text = q["number"]
        cells[1].text = q["type"]
        cells[2].text = "、".join(q["knowledge_points"])
        cells[3].text = f"{q['difficulty']['level']} / {q['difficulty']['label']}"
        cells[4].text = f"{q['suitability']['score']} / {q['suitability']['grade']} {q['suitability']['stage']}".strip()
        cells[5].text = "；".join(x for x in [q["difficulty"]["reason"], q["suitability"]["reason"]] if x)

    doc.add_page_break()
    doc.add_heading("四、变式题", level=1)
    for q in package["questions"]:
        if not q["variants"]:
            continue
        add_docx_paragraph(doc, f"原题 {q['number']} 变式：")
        for idx, variant in enumerate(q["variants"], start=1):
            add_docx_paragraph(doc, f"({idx}) {variant['stem']}")
            for option in variant["options"]:
                add_docx_paragraph(doc, option)
            if variant["answer"]:
                add_docx_paragraph(doc, f"答案：{variant['answer']}")
            if variant["explanation"]:
                add_docx_paragraph(doc, f"解析：{variant['explanation']}")

    doc.save(str(path))


def find_chinese_font() -> str:
    candidates = [
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    raise RuntimeError("找不到中文字体，无法生成 PDF")


def pdf_styles() -> dict[str, ParagraphStyle]:
    font_path = find_chinese_font()
    if "CNFont" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("CNFont", font_path))
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("CNTitle", parent=base["Title"], fontName="CNFont", fontSize=18, leading=24, spaceAfter=8),
        "h1": ParagraphStyle("CNH1", parent=base["Heading1"], fontName="CNFont", fontSize=14, leading=20, spaceBefore=10, spaceAfter=6),
        "normal": ParagraphStyle("CNNormal", parent=base["Normal"], fontName="CNFont", fontSize=10.5, leading=16),
        "small": ParagraphStyle("CNSmall", parent=base["Normal"], fontName="CNFont", fontSize=9, leading=13),
    }


def para(text: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(clean_text(text)), style)


def build_pdf(package: dict[str, Any], path: Path) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )
    story: list[Any] = []
    story.append(para(package["title"], styles["title"]))
    story.append(para(f"学科：{package['subject']}    年级：{package['grade']}    生成时间：{package['generated_at']}", styles["normal"]))
    story.append(para(f"来源质量：{package['source_quality']}", styles["normal"]))
    if package["processing_notes"]:
        story.append(para("处理说明：" + "；".join(package["processing_notes"]), styles["normal"]))

    story.append(para("一、清版试卷", styles["h1"]))
    for q in package["questions"]:
        story.append(para(f"{q['number']}. {q['clean_version'] or q['stem']}", styles["normal"]))
        for option in q["options"]:
            story.append(para(option, styles["small"]))
        story.append(Spacer(1, 4))

    story.append(PageBreak())
    story.append(para("二、答案与解析", styles["h1"]))
    for q in package["questions"]:
        story.append(para(f"{q['number']}. 答案：{q['answer'] or '未识别'}", styles["normal"]))
        if q["explanation"]:
            story.append(para(f"解析：{q['explanation']}", styles["small"]))
        story.append(Spacer(1, 4))

    story.append(para("三、题型、知识点、难度与适配度", styles["h1"]))
    rows: list[list[Any]] = [[para(x, styles["small"]) for x in ["题号", "题型", "知识点", "难度", "适配度"]]]
    for q in package["questions"]:
        rows.append(
            [
                para(q["number"], styles["small"]),
                para(q["type"], styles["small"]),
                para("、".join(q["knowledge_points"]), styles["small"]),
                para(f"{q['difficulty']['level']} / {q['difficulty']['label']}", styles["small"]),
                para(f"{q['suitability']['score']} / {q['suitability']['grade']} {q['suitability']['stage']}".strip(), styles["small"]),
            ]
        )
    table = Table(rows, colWidths=[18 * mm, 28 * mm, 58 * mm, 28 * mm, 42 * mm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "CNFont"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF7")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#888888")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)

    story.append(PageBreak())
    story.append(para("四、变式题", styles["h1"]))
    for q in package["questions"]:
        if not q["variants"]:
            continue
        story.append(para(f"原题 {q['number']} 变式：", styles["normal"]))
        for idx, variant in enumerate(q["variants"], start=1):
            story.append(para(f"({idx}) {variant['stem']}", styles["normal"]))
            for option in variant["options"]:
                story.append(para(option, styles["small"]))
            if variant["answer"]:
                story.append(para(f"答案：{variant['answer']}", styles["small"]))
            if variant["explanation"]:
                story.append(para(f"解析：{variant['explanation']}", styles["small"]))
            story.append(Spacer(1, 4))

    doc.build(story)


def build_markdown(package: dict[str, Any], path: Path) -> None:
    lines = [
        f"# {package['title']}",
        "",
        f"- 学科：{package['subject']}",
        f"- 年级：{package['grade']}",
        f"- 生成时间：{package['generated_at']}",
        f"- 来源质量：{package['source_quality']}",
        "",
        "## 题目分析",
        "",
    ]
    for q in package["questions"]:
        lines.extend(
            [
                f"### {q['number']}. {q['type']}",
                f"- 知识点：{'、'.join(q['knowledge_points'])}",
                f"- 难度：{q['difficulty']['level']} / {q['difficulty']['label']}，{q['difficulty']['reason']}",
                f"- 适配度：{q['suitability']['score']}，{q['suitability']['reason']}",
                f"- 答案：{q['answer']}",
                f"- 解析：{q['explanation']}",
                "",
            ]
        )
    path.write_text("\n".join(lines), encoding="utf-8")


def make_zip(files: list[Path], zip_path: Path, root: Path) -> None:
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for file_path in files:
            if file_path.exists():
                archive.write(file_path, file_path.relative_to(root))


def create_sample_package() -> dict[str, Any]:
    return normalize_package(
        {
            "title": "四年级英语旧试卷教学资产包",
            "subject": "英语",
            "grade": "四年级",
            "source_quality": "样例数据，用于测试排版输出",
            "processing_notes": ["这是本地自测样例，没有调用大模型"],
            "questions": [
                {
                    "number": "1",
                    "type": "选择题",
                    "stem": "Choose the right answer: I ____ a student.",
                    "options": ["A. am", "B. is", "C. are"],
                    "answer": "A",
                    "explanation": "I 后面用 am。",
                    "knowledge_points": ["be 动词", "主谓一致"],
                    "difficulty": {"level": 1, "label": "基础", "reason": "考查基础 be 动词搭配"},
                    "suitability": {"grade": "四年级", "stage": "同步练习", "score": 5, "reason": "适合课堂基础巩固"},
                    "variants": [
                        {
                            "stem": "Choose the right answer: She ____ my friend.",
                            "options": ["A. am", "B. is", "C. are"],
                            "answer": "B",
                            "explanation": "She 后面用 is。",
                            "knowledge_points": ["be 动词"],
                            "difficulty_level": 1,
                        }
                    ],
                }
            ],
        },
        title="",
        subject="",
        grade="",
    )


def generate_outputs(package: dict[str, Any], out_dir: Path) -> dict[str, str]:
    mkdir(out_dir)
    json_path = out_dir / "question_bank.json"
    md_path = out_dir / "analysis.md"
    docx_path = out_dir / "clean_paper.docx"
    pdf_path = out_dir / "clean_paper.pdf"
    zip_path = out_dir / "asset_package.zip"

    json_path.write_text(json.dumps(package, ensure_ascii=False, indent=2), encoding="utf-8")
    build_markdown(package, md_path)
    build_docx(package, docx_path)
    build_pdf(package, pdf_path)
    make_zip([json_path, md_path, docx_path, pdf_path], zip_path, out_dir)
    return {
        "question_bank_json": str(json_path),
        "analysis_md": str(md_path),
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        "zip": str(zip_path),
    }


def make_asset_package(
    source_path: str | Path,
    title: str = "",
    subject: str = "",
    grade: str = "",
    max_pages: int = 12,
) -> dict[str, Any]:
    source = Path(source_path).expanduser().resolve()
    if not source.exists() or not source.is_file():
        raise FileNotFoundError(f"source file not found: {source}")

    work_root = env_path("WORK_DIR", DEFAULT_WORK_DIR)
    output_root = env_path("OUTPUT_DIR", DEFAULT_OUTPUT_DIR)
    job_id = now_id()
    job_dir = mkdir(work_root / job_id)
    out_dir = mkdir(output_root / job_id)

    original_path = out_dir / ("original" + source.suffix.lower())
    shutil.copy2(source, original_path)

    prepared = prepare_source(source, job_dir, max_pages=max_pages)
    raw = analyze_with_llm(prepared, title=title, subject=subject, grade=grade)
    package = normalize_package(raw, title=title, subject=subject, grade=grade)
    outputs = generate_outputs(package, out_dir)
    make_zip([Path(v) for v in outputs.values() if str(v).endswith((".json", ".md", ".docx", ".pdf"))] + [original_path], out_dir / "asset_package.zip", out_dir)
    outputs["zip"] = str(out_dir / "asset_package.zip")
    return {
        "ok": True,
        "job_id": job_id,
        "source": str(source),
        "out_dir": str(out_dir),
        "question_count": len(package["questions"]),
        "outputs": outputs,
    }


def self_test() -> dict[str, Any]:
    output_root = env_path("OUTPUT_DIR", DEFAULT_OUTPUT_DIR)
    job_id = "selftest_" + now_id()
    out_dir = mkdir(output_root / job_id)
    package = create_sample_package()
    outputs = generate_outputs(package, out_dir)
    return {"ok": True, "job_id": job_id, "out_dir": str(out_dir), "outputs": outputs}


def main() -> int:
    parser = argparse.ArgumentParser(description="旧试卷教学资产包生成器")
    parser.add_argument("source", nargs="?", help="PDF、图片、docx 或文本文件")
    parser.add_argument("--title", default="")
    parser.add_argument("--subject", default="")
    parser.add_argument("--grade", default="")
    parser.add_argument("--max-pages", type=int, default=12)
    parser.add_argument("--self-test", action="store_true")
    args = parser.parse_args()

    if args.self_test:
        print(json.dumps(self_test(), ensure_ascii=False, indent=2))
        return 0
    if not args.source:
        parser.error("source is required unless --self-test is used")

    result = make_asset_package(args.source, title=args.title, subject=args.subject, grade=args.grade, max_pages=args.max_pages)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
